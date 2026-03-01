"""
Generate Rohit Ranvir's master resume as a formatted .docx file.
Run: python build_resume.py  (from D:\Project\job-agent)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "output", "Rohit_Ranvir_Resume.docx")

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Helpers ──
def add_section_heading(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x54, 0x7D)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F547D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(text, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_role_header(title, org, location, date_range):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tcBorders.append(el)
        tcPr.append(tcBorders)
    left  = table.cell(0, 0)
    right = table.cell(0, 1)

    lp1 = left.paragraphs[0]
    lp1.paragraph_format.space_before = Pt(4)
    lp1.paragraph_format.space_after  = Pt(0)
    r = lp1.add_run(title)
    r.bold = True; r.font.size = Pt(10)

    lp2 = left.add_paragraph()
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(1)
    r2 = lp2.add_run(org)
    r2.italic = True; r2.font.size = Pt(9.5)
    lp2.add_run(f"  —  {location}").font.size = Pt(9.5)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after  = Pt(1)
    rd = rp.add_run(date_range)
    rd.font.size = Pt(9.5)

def add_project_header(title, subtitle, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(title); r1.bold = True; r1.font.size = Pt(10)
    if subtitle:
        p.add_run(f"  —  {subtitle}").font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {tech}")
    r3.italic = True; r3.font.size = Pt(9)

# ═══ HEADER ═══
hp = doc.add_paragraph()
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.space_before = Pt(0)
hp.paragraph_format.space_after  = Pt(2)
hr = hp.add_run("ROHIT RANVIR")
hr.bold = True; hr.font.size = Pt(18)
hr.font.color.rgb = RGBColor(0x1F, 0x54, 0x7D)

cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(0)
cp.paragraph_format.space_after  = Pt(4)
cp.add_run("Maharashtra, India  |  +91-9158000676  |  rohitranveer358@gmail.com  |  LinkedIn  |  GitHub").font.size = Pt(9.5)

# ═══ PROFESSIONAL SUMMARY ═══
add_section_heading("Professional Summary")
sp = doc.add_paragraph()
sp.paragraph_format.space_before = Pt(2); sp.paragraph_format.space_after = Pt(2)
sp.add_run(
    "Results-driven Data Analyst skilled in building end-to-end ETL pipelines, ML models, and interactive BI dashboards "
    "using Python, SQL, and Power BI. Proven ability to transform large datasets into actionable insights and automate "
    "workflows at scale. Eager to apply data engineering and analytical expertise to high-impact, large-scale problems."
).font.size = Pt(9.5)

# ═══ TECHNICAL SKILLS ═══
add_section_heading("Technical Skills")
skills = [
    ("Languages & Query",   "Python (Pandas, NumPy, Scikit-learn, Matplotlib, Seaborn), SQL (MySQL, SQLite, SQL Server)"),
    ("Data & ML",           "Data Cleaning, EDA, Feature Engineering, Statistical Analysis, Machine Learning, NLP, Time-Series Forecasting"),
    ("BI & Visualization",  "Power BI (DAX, Power Query), Streamlit, Plotly, MS Excel (Pivot Tables, XLOOKUP)"),
    ("Engineering & Tools", "ETL Pipeline Design, REST APIs, GitHub Actions (CI/CD), Web Scraping (Playwright), Google Sheets"),
]
for label, val in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label + ": "); r1.bold = True; r1.font.size = Pt(9.5)
    p.add_run(val).font.size = Pt(9.5)

# ═══ EXPERIENCE ═══
add_section_heading("Experience")
add_role_header("Data Analyst Intern", "Codesoft Private Limited", "Maharashtra, India", "02/2026 – Present")
add_bullet("Developed interactive Power BI dashboards tracking 10+ KPIs, improving stakeholder reporting efficiency by 40% and enabling real-time data-driven decision-making.")
add_bullet("Executed EDA on large datasets to surface actionable business trends, directly informing product and operations strategy.")
add_bullet("Automated routine reporting workflows using Python scripts, eliminating 100% of manual data-collection effort and saving 5+ analyst hours per week.")

# ═══ PROJECTS ═══
add_section_heading("Projects")

add_project_header("Inflation Basket Tracker", "Automated Price Monitoring & ML Forecasting", "Python, Playwright, SQLite, Scikit-learn, Streamlit, GitHub Actions")
add_bullet("Engineered an end-to-end automated ETL pipeline that scrapes daily grocery prices for 5+ SKUs using Playwright, storing structured results in SQLite; automated via GitHub Actions cron jobs, eliminating 100% of manual collection effort.")
add_bullet("Applied time-series feature engineering (date ordinal, seasonality flags, day-of-week) and trained a Random Forest regression model achieving ~76% R² accuracy to produce 7-day grocery cost forecasts.")
add_bullet("Deployed an interactive Streamlit dashboard with Plotly visualizations for historical price trends, daily inflation rates, and ML-generated forecasts.")

add_project_header("Wikipedia Edit War Analyzer", "NLP & Controversy Detection", "Python, Pandas, Plotly, Streamlit, MediaWiki REST API, NLP")
add_bullet("Built a data analysis system processing up to 5,000 Wikipedia revision records via the MediaWiki REST API; developed a weighted Controversy Score algorithm (0–100) using revert rate, editor density, and edit-spike frequency with log-normalization.")
add_bullet("Implemented keyword-based NLP classification of edit comments (hostility, dispute, revert, protection) and surfaced insights in an interactive Streamlit dashboard with 6 Plotly charts and one-click CSV export.")

add_project_header("Power BI Business Performance Dashboard", "", "Power BI, Power Query, DAX, Excel")
add_bullet("Designed an interactive Power BI dashboard with dynamic filters and slicers, reducing manual reporting effort and improving stakeholder reporting efficiency by 40%.")
add_bullet("Engineered DAX-based calculated measures and KPIs to track real-time trends and category-level performance across multiple integrated data sources.")

add_project_header("Customer Churn Analysis", "", "Python, Pandas, Seaborn, Matplotlib, EDA")
add_bullet("Performed EDA on customer demographic and behavioural data to identify churn patterns across tenure, service usage, and billing behaviour.")
add_bullet("Conducted statistical correlation analysis to uncover key churn-driving features and segment high-risk customer groups, enabling targeted business retention strategies.")

# ═══ EDUCATION ═══
add_section_heading("Education")
add_role_header("B.E. in Computer Science and Engineering", "Babasaheb Naik College of Engineering, Pusad", "Amravati University, Maharashtra, India", "08/2021 – 05/2025")
ep = doc.add_paragraph()
ep.paragraph_format.space_before = Pt(1); ep.paragraph_format.space_after = Pt(1)
ep.add_run("Coursework: DBMS, Data Structures & Algorithms, Statistics, Operating Systems, Computer Networks").font.size = Pt(9.5)

# ═══ CERTIFICATIONS ═══
add_section_heading("Certifications")
certp = doc.add_paragraph()
certp.paragraph_format.space_before = Pt(2); certp.paragraph_format.space_after = Pt(0)
rc = certp.add_run("Data Analyst — Simplilearn  (2025)")
rc.bold = True; rc.font.size = Pt(9.5)
certp2 = doc.add_paragraph()
certp2.paragraph_format.space_before = Pt(0); certp2.paragraph_format.space_after = Pt(2)
certp2.add_run("Completed professional certification covering data analysis, SQL, Python, data visualization, and business intelligence fundamentals.").font.size = Pt(9.5)

# ── Save ──
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Resume saved to: {OUTPUT_PATH}")
