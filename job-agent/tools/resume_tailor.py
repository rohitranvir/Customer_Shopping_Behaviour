"""
resume_tailor.py — Uses Claude AI to tailor the master resume for a specific job.
"""
import os
import re
import json
import shutil
from pathlib import Path
from dotenv import load_dotenv

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv()

OUTPUT_DIR = Path(__file__).parent.parent / "output" / "tailored_resumes"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _safe_filename(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", text)[:40]


def _get_resume_text(resume_path: Path) -> str:
    """Extract plain text from a .docx resume."""
    if not DOCX_AVAILABLE:
        return "[python-docx not installed — install it to read resume]"
    if not resume_path.exists():
        return "[master_resume.docx not found — please add your resume]"
    doc = Document(str(resume_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _call_claude(prompt: str) -> str:
    """Send a prompt to Claude and return the response text."""
    if not ANTHROPIC_AVAILABLE:
        raise RuntimeError("anthropic package not installed. Run: pip install anthropic")
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key or api_key.startswith("your_"):
        raise RuntimeError("ANTHROPIC_API_KEY not set in .env file")
    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def tailor_resume(
    company: str,
    role: str,
    job_description: str,
    profile: dict,
    master_resume_path: Path,
) -> tuple[float, str]:
    """
    Tailors the master resume for the given job.

    Returns:
        (match_score: float, output_path: str)
        match_score is 0–100; output_path is the saved .docx path.
    """
    resume_text = _get_resume_text(master_resume_path)

    prompt = f"""You are an expert ATS resume optimizer. 

CANDIDATE PROFILE:
Name: {profile['name']}
Skills: {', '.join(profile['skills'])}
Experience: {profile['experience_level']}

ORIGINAL RESUME:
{resume_text}

JOB DETAILS:
Company: {company}
Role: {role}
Description:
{job_description[:3000]}

TASKS:
1. Calculate a match score (0-100) based on how well the candidate's skills and experience match the job requirements.
2. Rewrite the resume summary section (2-3 sentences) to highlight the most relevant skills for THIS specific job.
3. Suggest 3-5 bullet points for a projects/skills section that best match the job requirements.

Respond in this exact JSON format:
{{
  "match_score": 75,
  "tailored_summary": "...",
  "tailored_bullets": ["...", "...", "..."],
  "keywords_added": ["keyword1", "keyword2"]
}}"""

    try:
        response_text = _call_claude(prompt)
        # Extract JSON from the response
        json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
        if not json_match:
            raise ValueError("No JSON found in Claude response")
        data = json.loads(json_match.group())
        match_score = float(data.get("match_score", 50))
        tailored_summary = data.get("tailored_summary", "")
        tailored_bullets = data.get("tailored_bullets", [])
    except Exception as e:
        print(f"  [WARN] Claude tailor failed: {e} — using original resume with score=60")
        match_score = 60.0
        tailored_summary = ""
        tailored_bullets = []

    # Save the tailored resume
    safe_company = _safe_filename(company)
    safe_role = _safe_filename(role)
    output_path = OUTPUT_DIR / f"{safe_company}_{safe_role}.docx"

    if DOCX_AVAILABLE and master_resume_path.exists():
        doc = Document(str(master_resume_path))
        # Insert a tailored note at the top if we have one
        if tailored_summary:
            doc.paragraphs[0].insert_paragraph_before(
                f"[TAILORED FOR: {company} — {role}]\n{tailored_summary}"
            )
        if tailored_bullets:
            last_para = doc.add_paragraph("Key Highlights for this Role:")
            for bullet in tailored_bullets:
                doc.add_paragraph(f"• {bullet}")
        doc.save(str(output_path))
    else:
        # Fallback: copy original or create placeholder
        if master_resume_path.exists():
            shutil.copy(str(master_resume_path), str(output_path))
        else:
            output_path.write_text(
                f"Tailored resume for {role} at {company}\n"
                f"Candidate: {profile['name']}\n"
                f"Summary: {tailored_summary}\n"
                f"Bullets: {chr(10).join(tailored_bullets)}"
            )

    return match_score, str(output_path)
